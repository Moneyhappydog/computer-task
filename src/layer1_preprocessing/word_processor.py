"""
Word 文档处理器
负责将 .docx 文档转换为结构化 Markdown
"""
from pathlib import Path
from typing import Dict, Any
import logging
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import re

logger = logging.getLogger(__name__)

class WordProcessor:
    """Word文档处理器"""
    
    def __init__(self):
        """初始化Word处理器"""
        self.supported_formats = ['.docx']
        logger.info("✅ Word处理器初始化完成")
    
    def process(self, file_path: Path) -> Dict[str, Any]:
        """
        处理Word文档
        
        Args:
            file_path: Word文件路径
            
        Returns:
            包含Markdown内容和元数据的字典
        """
        logger.info(f"📄 开始处理Word文档: {file_path.name}")
        
        try:
            # 读取Word文档
            doc = Document(file_path)
            
            # 提取内容
            markdown_lines = []
            metadata = {
                'file_name': file_path.name,
                'file_type': 'docx',
                'paragraphs': 0,
                'tables': 0,
                'images': 0,
                'headings': {}
            }
            
            # 处理段落
            for para in doc.paragraphs:
                md_text = self._process_paragraph(para, metadata)
                if md_text:
                    markdown_lines.append(md_text)
                    metadata['paragraphs'] += 1
            
            # 处理表格
            for table in doc.tables:
                md_table = self._process_table(table)
                if md_table:
                    markdown_lines.append(md_table)
                    metadata['tables'] += 1
            
            # 统计图片
            metadata['images'] = len(doc.inline_shapes)
            
            markdown_content = '\n\n'.join(markdown_lines)
            
            logger.info(f"✅ Word处理完成: {metadata['paragraphs']}段落, "
                       f"{metadata['tables']}表格, {metadata['images']}图片")
            
            return {
                'markdown': markdown_content,
                'metadata': metadata,
                'success': True
            }
            
        except Exception as e:
            logger.error(f"❌ Word处理失败: {e}")
            return {
                'markdown': '',
                'metadata': {},
                'success': False,
                'error': str(e)
            }
    
    def _process_paragraph(self, para, metadata: Dict) -> str:
        """
        处理单个段落
        
        Args:
            para: Word段落对象
            metadata: 元数据字典
            
        Returns:
            Markdown格式的段落文本
        """
        text = para.text.strip()
        if not text:
            return ""
        
        # 检测标题
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            
            # 记录标题层级
            if f'h{level}' not in metadata['headings']:
                metadata['headings'][f'h{level}'] = 0
            metadata['headings'][f'h{level}'] += 1
            
            return f"{'#' * level} {text}"
        
        # 检测列表
        if para.style.name.startswith('List'):
            if 'Bullet' in para.style.name:
                return f"- {text}"
            elif 'Number' in para.style.name:
                return f"1. {text}"
        
        # 检测代码块（使用等宽字体）
        if para.runs and len(para.runs) > 0:
            first_run = para.runs[0]
            if first_run.font.name in ['Courier New', 'Consolas', 'Courier']:
                return f"```\n{text}\n```"
        
        # 检测加粗/斜体
        formatted_text = self._apply_inline_formatting(para)
        
        return formatted_text
    
    def _apply_inline_formatting(self, para) -> str:
        """
        应用行内格式（加粗、斜体等）
        
        Args:
            para: Word段落对象
            
        Returns:
            带Markdown格式的文本
        """
        result = []
        
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # 加粗
            if run.bold:
                text = f"**{text}**"
            
            # 斜体
            if run.italic:
                text = f"*{text}*"
            
            # 代码（等宽字体）
            if run.font.name in ['Courier New', 'Consolas', 'Courier']:
                text = f"`{text}`"
            
            result.append(text)
        
        return ''.join(result)
    
    def _process_table(self, table) -> str:
        """
        处理Word表格
        
        Args:
            table: Word表格对象
            
        Returns:
            Markdown格式的表格
        """
        if not table.rows:
            return ""
        
        markdown_lines = []
        
        # 表头
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        markdown_lines.append('| ' + ' | '.join(header_cells) + ' |')
        
        # 分隔线
        markdown_lines.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')
        
        # 数据行
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            markdown_lines.append('| ' + ' | '.join(cells) + ' |')
        
        return '\n'.join(markdown_lines)
    
    def is_supported(self, file_path: Path) -> bool:
        """
        检查文件格式是否支持
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否支持该格式
        """
        return file_path.suffix.lower() in self.supported_formats


# 测试代码
if __name__ == "__main__":
    from src.utils.logger import setup_logger
    
    setup_logger("word_processor")
    
    processor = WordProcessor()
    
    # 测试文件
    test_file = Path("data/input/test_document.docx")
    
    if test_file.exists():
        result = processor.process(test_file)
        
        if result['success']:
            print("\n" + "="*70)
            print("Markdown 内容:")
            print("="*70)
            print(result['markdown'][:500])
            
            print("\n" + "="*70)
            print("元数据:")
            print("="*70)
            print(result['metadata'])
    else:
        print(f"❌ 测试文件不存在: {test_file}")